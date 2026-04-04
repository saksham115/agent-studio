"""Document parsing and chunking for knowledge-base ingestion.

Supports PDF (via PyMuPDF), DOCX (via python-docx), XLSX (via openpyxl),
plain text, and CSV files.  Documents are split into overlapping chunks
suitable for embedding.
"""

from __future__ import annotations

import csv
import io
import logging
import math
from dataclasses import dataclass

logger = logging.getLogger(__name__)


@dataclass
class TextChunk:
    """A single chunk of text extracted from a document."""

    content: str
    index: int
    token_count: int


def _estimate_tokens(text: str) -> int:
    """Rough token estimate: word count * 1.3, minimum 1."""
    if not text:
        return 0
    return max(1, math.ceil(len(text.split()) * 1.3))


class DocumentProcessor:
    """Parses documents and splits the resulting text into chunks."""

    # ------------------------------------------------------------------
    # Public API
    # ------------------------------------------------------------------

    async def process_document(
        self,
        file_content: bytes,
        filename: str,
        source_type: str,
    ) -> list[TextChunk]:
        """Parse *file_content* according to *source_type* and return chunks.

        Raises ``ValueError`` if the source type is unsupported or the
        document contains no extractable text.
        """
        source_type = source_type.lower()

        if source_type == "pdf":
            raw_text = self._parse_pdf(file_content)
        elif source_type in ("txt", "text"):
            raw_text = self._parse_txt(file_content)
        elif source_type == "csv":
            raw_text = self._parse_csv(file_content)
        elif source_type == "docx":
            raw_text = self._parse_docx(file_content)
        elif source_type in ("xlsx", "xls"):
            raw_text = self._parse_xlsx(file_content)
        else:
            raise ValueError(f"Unsupported source type: {source_type}")

        if not raw_text or not raw_text.strip():
            raise ValueError(
                f"No text could be extracted from '{filename}' (type={source_type})"
            )

        return self._chunk_text(raw_text)

    # ------------------------------------------------------------------
    # Parsers
    # ------------------------------------------------------------------

    def _parse_pdf(self, content: bytes) -> str:
        """Extract text from a PDF using PyMuPDF (fitz)."""
        try:
            import fitz  # PyMuPDF
        except ImportError as exc:
            raise RuntimeError(
                "PyMuPDF is required for PDF parsing. Install it with: pip install PyMuPDF"
            ) from exc

        pages: list[str] = []
        with fitz.open(stream=content, filetype="pdf") as doc:
            for page_num, page in enumerate(doc):
                text = page.get_text("text")
                if text and text.strip():
                    pages.append(text.strip())
                else:
                    logger.debug("PDF page %d produced no text", page_num)
        return "\n\n".join(pages)

    def _parse_txt(self, content: bytes) -> str:
        """Decode raw bytes to text (UTF-8 with latin-1 fallback)."""
        try:
            return content.decode("utf-8")
        except UnicodeDecodeError:
            return content.decode("latin-1")

    def _parse_csv(self, content: bytes) -> str:
        """Parse CSV into human-readable text rows.

        Each row is rendered as ``header1: value1 | header2: value2 | ...``
        so that the resulting text is easier for an LLM to reason about.
        """
        try:
            text = content.decode("utf-8")
        except UnicodeDecodeError:
            text = content.decode("latin-1")

        reader = csv.reader(io.StringIO(text))
        rows = list(reader)

        if not rows:
            return ""

        headers = rows[0]
        output_lines: list[str] = []

        for row in rows[1:]:
            parts: list[str] = []
            for idx, value in enumerate(row):
                header = headers[idx] if idx < len(headers) else f"col_{idx}"
                parts.append(f"{header}: {value}")
            output_lines.append(" | ".join(parts))

        return "\n".join(output_lines)

    def _parse_docx(self, content: bytes) -> str:
        """Extract text from a DOCX file using python-docx."""
        try:
            from docx import Document
        except ImportError as exc:
            raise RuntimeError(
                "python-docx is required for DOCX parsing. Install it with: pip install python-docx"
            ) from exc

        doc = Document(io.BytesIO(content))
        paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    paragraphs.append(" | ".join(cells))

        return "\n\n".join(paragraphs)

    def _parse_xlsx(self, content: bytes) -> str:
        """Extract text from an XLSX file using openpyxl."""
        try:
            from openpyxl import load_workbook
        except ImportError as exc:
            raise RuntimeError(
                "openpyxl is required for XLSX parsing. Install it with: pip install openpyxl"
            ) from exc

        wb = load_workbook(io.BytesIO(content), read_only=True, data_only=True)
        output_lines: list[str] = []

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            output_lines.append(f"## Sheet: {sheet_name}")
            headers: list[str] = []
            for row_idx, row in enumerate(ws.iter_rows(values_only=True)):
                values = [str(v) if v is not None else "" for v in row]
                if not any(values):
                    continue
                if row_idx == 0:
                    headers = values
                    output_lines.append(" | ".join(headers))
                elif headers:
                    parts = [f"{h}: {v}" for h, v in zip(headers, values) if v]
                    if parts:
                        output_lines.append(" | ".join(parts))
                else:
                    output_lines.append(" | ".join(v for v in values if v))

        wb.close()
        return "\n".join(output_lines)

    # ------------------------------------------------------------------
    # Chunking
    # ------------------------------------------------------------------

    def _chunk_text(
        self,
        text: str,
        chunk_size: int = 500,
        overlap: int = 50,
    ) -> list[TextChunk]:
        """Recursive character splitter.

        Strategy:
        1. Split on double-newlines (paragraphs).
        2. If a paragraph is still too large, split on sentence boundaries.
        3. If a sentence is still too large, split on whitespace (words).

        *chunk_size* and *overlap* are expressed in estimated tokens.
        """
        paragraphs = self._split_paragraphs(text)

        # Flatten paragraphs that are too large into smaller pieces.
        fragments: list[str] = []
        for para in paragraphs:
            if _estimate_tokens(para) <= chunk_size:
                fragments.append(para)
            else:
                fragments.extend(self._split_sentences(para, chunk_size))

        # Merge small fragments back into chunks, respecting the size limit,
        # and add the requested overlap.
        chunks: list[TextChunk] = []
        current_parts: list[str] = []
        current_tokens = 0

        for frag in fragments:
            frag_tokens = _estimate_tokens(frag)
            if current_tokens + frag_tokens > chunk_size and current_parts:
                # Emit chunk
                chunk_text = "\n\n".join(current_parts).strip()
                if chunk_text:
                    chunks.append(
                        TextChunk(
                            content=chunk_text,
                            index=len(chunks),
                            token_count=_estimate_tokens(chunk_text),
                        )
                    )
                # Overlap: keep trailing parts whose token total <= overlap
                overlap_parts: list[str] = []
                overlap_tokens = 0
                for part in reversed(current_parts):
                    part_tokens = _estimate_tokens(part)
                    if overlap_tokens + part_tokens > overlap:
                        break
                    overlap_parts.insert(0, part)
                    overlap_tokens += part_tokens
                current_parts = overlap_parts
                current_tokens = overlap_tokens

            current_parts.append(frag)
            current_tokens += frag_tokens

        # Flush remaining
        if current_parts:
            chunk_text = "\n\n".join(current_parts).strip()
            if chunk_text:
                chunks.append(
                    TextChunk(
                        content=chunk_text,
                        index=len(chunks),
                        token_count=_estimate_tokens(chunk_text),
                    )
                )

        return chunks

    # ------------------------------------------------------------------
    # Splitting helpers
    # ------------------------------------------------------------------

    @staticmethod
    def _split_paragraphs(text: str) -> list[str]:
        """Split text on paragraph boundaries (double newline)."""
        parts = text.split("\n\n")
        return [p.strip() for p in parts if p.strip()]

    @staticmethod
    def _split_sentences(text: str, max_tokens: int) -> list[str]:
        """Split text on sentence boundaries; fall back to word-level."""
        # Simple sentence boundary heuristic: split on '. ', '! ', '? '
        import re

        sentences = re.split(r"(?<=[.!?])\s+", text)
        result: list[str] = []
        buf: list[str] = []
        buf_tokens = 0

        for sent in sentences:
            sent_tokens = _estimate_tokens(sent)

            # Single sentence larger than the limit -- split on words
            if sent_tokens > max_tokens:
                if buf:
                    result.append(" ".join(buf))
                    buf = []
                    buf_tokens = 0
                result.extend(DocumentProcessor._split_words(sent, max_tokens))
                continue

            if buf_tokens + sent_tokens > max_tokens and buf:
                result.append(" ".join(buf))
                buf = []
                buf_tokens = 0

            buf.append(sent)
            buf_tokens += sent_tokens

        if buf:
            result.append(" ".join(buf))

        return result

    @staticmethod
    def _split_words(text: str, max_tokens: int) -> list[str]:
        """Last-resort split: break on whitespace."""
        words = text.split()
        result: list[str] = []
        current: list[str] = []
        current_tokens = 0

        for word in words:
            word_tokens = _estimate_tokens(word)
            if current_tokens + word_tokens > max_tokens and current:
                result.append(" ".join(current))
                current = []
                current_tokens = 0
            current.append(word)
            current_tokens += word_tokens

        if current:
            result.append(" ".join(current))

        return result
